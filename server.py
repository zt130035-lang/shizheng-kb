#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
考公时政知识库系统 - 后端服务
启动: python server.py
访问: http://localhost:5678

目录结构:
  shizheng_kb/
  ├── server.py          # 本文件
  ├── static/            # 前端静态文件
  │   └── index.html
  └── data/
      ├── news_archive/  # 爬取的新闻 markdown
      ├── daily_reports/ # 日报/周报
      ├── knowledge_db/  # ChromaDB 向量库
      └── pdf_uploads/   # 上传的 PDF 文件
"""

import os
import re
import json
import glob
import hashlib
import fitz  # PyMuPDF
import requests
import chromadb
import markdown
from flask import Flask, request, jsonify, send_from_directory
from datetime import datetime
from typing import List
from werkzeug.utils import secure_filename

# ========== 路径配置 ==========
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# 云端部署时使用 /data 持久化目录，本地开发使用项目内 data/
DATA_DIR = os.environ.get("DATA_DIR", os.path.join(BASE_DIR, "data"))
NEWS_ARCHIVE_DIR = os.path.join(DATA_DIR, "news_archive")
DAILY_REPORTS_DIR = os.path.join(DATA_DIR, "daily_reports")
KNOWLEDGE_DB_DIR = os.path.join(DATA_DIR, "knowledge_db")
PDF_UPLOADS_DIR = os.path.join(DATA_DIR, "pdf_uploads")
STATIC_DIR = os.path.join(BASE_DIR, "static")

# 确保目录存在
for d in [NEWS_ARCHIVE_DIR, DAILY_REPORTS_DIR, KNOWLEDGE_DB_DIR, PDF_UPLOADS_DIR]:
    os.makedirs(d, exist_ok=True)

# 用户数据文件
USER_DATA_FILE = os.path.join(DATA_DIR, "user_data.json")

# ========== API 配置 ==========
SILICONFLOW_API_KEY = os.environ.get("SILICONFLOW_API_KEY", "sk-bapmemmmswmycyamlywctomymaievhzmlyznvenqosetlgxa")
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-97ca5455764543a8b57f63f3f9cacfef")
EMBEDDING_MODEL = "BAAI/bge-m3"
EMBEDDING_API_URL = "https://api.siliconflow.cn/v1/embeddings"

# GitHub 自动同步配置
GITHUB_TOKEN = os.environ.get("GITHUB_TOKEN", "")
GITHUB_REPO = os.environ.get("GITHUB_REPO", "zt130035-lang/shizheng-kb")

app = Flask(__name__, static_folder=STATIC_DIR)

# ========== 工具函数 ==========

def git_sync_to_github(message: str = "auto: sync data"):
    """将数据文件自动同步到GitHub（通过API，无需本地git）"""
    if not GITHUB_TOKEN:
        print("[GIT-SYNC] 未配置 GITHUB_TOKEN，跳过同步")
        return False

    import base64
    api_base = f"https://api.github.com/repos/{GITHUB_REPO}"
    headers = {
        "Authorization": f"token {GITHUB_TOKEN}",
        "Accept": "application/vnd.github.v3+json"
    }

    # 收集需要同步的文件
    sync_files = []

    # 新闻归档
    for f in glob.glob(os.path.join(NEWS_ARCHIVE_DIR, "*.md")):
        rel_path = f"data/news_archive/{os.path.basename(f)}"
        sync_files.append((f, rel_path))

    # 日报
    for f in glob.glob(os.path.join(DAILY_REPORTS_DIR, "*.md")):
        rel_path = f"data/daily_reports/{os.path.basename(f)}"
        sync_files.append((f, rel_path))

    # 题库 JSON
    pdf_q_dir = os.path.join(DATA_DIR, "pdf_questions")
    if os.path.isdir(pdf_q_dir):
        for f in glob.glob(os.path.join(pdf_q_dir, "*.json")):
            rel_path = f"data/pdf_questions/{os.path.basename(f)}"
            sync_files.append((f, rel_path))

    if not sync_files:
        print("[GIT-SYNC] 无文件需要同步")
        return True

    # 获取最新 commit SHA（main 分支）
    try:
        ref_resp = requests.get(f"{api_base}/git/ref/heads/main", headers=headers, timeout=15)
        if ref_resp.status_code != 200:
            print(f"[GIT-SYNC] 获取ref失败: {ref_resp.status_code}")
            return False
        latest_commit_sha = ref_resp.json()["object"]["sha"]

        # 获取该 commit 的 tree
        commit_resp = requests.get(f"{api_base}/git/commits/{latest_commit_sha}", headers=headers, timeout=15)
        base_tree_sha = commit_resp.json()["tree"]["sha"]
    except Exception as e:
        print(f"[GIT-SYNC] 获取仓库信息失败: {e}")
        return False

    # 创建 blobs 并构建 tree
    tree_items = []
    uploaded = 0
    for local_path, repo_path in sync_files:
        try:
            with open(local_path, "rb") as f:
                content = f.read()
            b64_content = base64.b64encode(content).decode("utf-8")
            blob_resp = requests.post(
                f"{api_base}/git/blobs",
                headers=headers,
                json={"content": b64_content, "encoding": "base64"},
                timeout=30
            )
            if blob_resp.status_code == 201:
                blob_sha = blob_resp.json()["sha"]
                tree_items.append({
                    "path": repo_path,
                    "mode": "100644",
                    "type": "blob",
                    "sha": blob_sha
                })
                uploaded += 1
        except Exception as e:
            print(f"[GIT-SYNC] 上传blob失败 {repo_path}: {e}")

    if not tree_items:
        print("[GIT-SYNC] 无文件上传成功")
        return False

    # 创建新 tree
    try:
        tree_resp = requests.post(
            f"{api_base}/git/trees",
            headers=headers,
            json={"base_tree": base_tree_sha, "tree": tree_items},
            timeout=30
        )
        if tree_resp.status_code != 201:
            print(f"[GIT-SYNC] 创建tree失败: {tree_resp.text[:200]}")
            return False
        new_tree_sha = tree_resp.json()["sha"]
    except Exception as e:
        print(f"[GIT-SYNC] 创建tree异常: {e}")
        return False

    # 创建 commit
    try:
        commit_data = {
            "message": message,
            "tree": new_tree_sha,
            "parents": [latest_commit_sha]
        }
        commit_resp = requests.post(
            f"{api_base}/git/commits",
            headers=headers,
            json=commit_data,
            timeout=15
        )
        if commit_resp.status_code != 201:
            print(f"[GIT-SYNC] 创建commit失败: {commit_resp.text[:200]}")
            return False
        new_commit_sha = commit_resp.json()["sha"]
    except Exception as e:
        print(f"[GIT-SYNC] 创建commit异常: {e}")
        return False

    # 更新 ref 指向新 commit
    try:
        ref_update = requests.patch(
            f"{api_base}/git/refs/heads/main",
            headers=headers,
            json={"sha": new_commit_sha},
            timeout=15
        )
        if ref_update.status_code == 200:
            print(f"[GIT-SYNC] ✅ 同步成功！{uploaded}个文件已推送到GitHub")
            return True
        else:
            print(f"[GIT-SYNC] 更新ref失败: {ref_update.text[:200]}")
            return False
    except Exception as e:
        print(f"[GIT-SYNC] 更新ref异常: {e}")
        return False

def get_embedding(texts: List[str]) -> List[List[float]]:
    """调用SiliconFlow获取文本向量"""
    headers = {
        "Authorization": f"Bearer {SILICONFLOW_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {"model": EMBEDDING_MODEL, "input": texts, "encoding_format": "float"}
    try:
        resp = requests.post(EMBEDDING_API_URL, headers=headers, json=payload, timeout=30)
        resp.raise_for_status()
        return [item["embedding"] for item in resp.json()["data"]]
    except Exception as e:
        print(f"[Embedding Error] {e}")
        return []


def get_chroma_client():
    return chromadb.PersistentClient(path=KNOWLEDGE_DB_DIR)


def call_deepseek(prompt: str, system: str = "") -> str:
    """调用DeepSeek AI"""
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})
    payload = {"model": "deepseek-chat", "messages": messages, "temperature": 0.3}
    try:
        resp = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers=headers, json=payload, timeout=60
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"AI调用失败: {e}"


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
    """将长文本分块"""
    if len(text) <= chunk_size:
        return [text] if text.strip() else []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def extract_pdf_text(filepath: str) -> str:
    """从PDF提取全部文本"""
    doc = fitz.open(filepath)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def extract_docx_text(filepath: str) -> str:
    """从Word(.docx)文档提取全部文本，含表格"""
    from docx import Document
    document = Document(filepath)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_document_text(filepath: str, filename: str = "") -> str:
    """根据扩展名自动提取文档文本，支持 PDF / Word(.docx) / 纯文本"""
    name = (filename or filepath).lower()
    if name.endswith(".pdf"):
        return extract_pdf_text(filepath)
    if name.endswith(".docx"):
        return extract_docx_text(filepath)
    if name.endswith(".doc"):
        raise ValueError("暂不支持旧版 .doc 格式，请另存为 .docx 或 PDF 后上传")
    if name.endswith((".txt", ".md")):
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    # 兜底：按文本读取
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _parse_json_loose(text: str):
    """从AI返回中宽松解析JSON对象，兼容代码块包裹/前后多余文本"""
    if not text:
        return None
    t = text.strip()
    # 去掉markdown代码块标记
    t = re.sub(r"^```(?:json)?\s*", "", t)
    t = re.sub(r"\s*```$", "", t).strip()
    try:
        return json.loads(t)
    except Exception:
        pass
    # 提取第一个 {...} 块
    match = re.search(r"\{.*\}", t, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            return None
    return None


def sanitize_collection_name(name: str) -> str:
    """将用户输入的名称转为ChromaDB合法集合名
    ChromaDB要求: 3-512字符, [a-zA-Z0-9._-], 首尾必须是[a-zA-Z0-9]
    """
    # 如果已经是合法名称直接返回
    if re.match(r'^[a-zA-Z0-9][a-zA-Z0-9._-]{1,510}[a-zA-Z0-9]$', name):
        return name
    # 否则用 kb_ + hash 前8位作为ID
    h = hashlib.md5(name.encode()).hexdigest()[:8]
    return f"kb_{h}"


def load_user_data() -> dict:
    """加载用户数据"""
    if os.path.exists(USER_DATA_FILE):
        with open(USER_DATA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"wrong_questions": [], "quiz_history": [], "favorites": []}


def save_user_data(data: dict):
    """保存用户数据"""
    with open(USER_DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# ========== API 路由 ==========

@app.route("/")
def index():
    return send_from_directory(STATIC_DIR, "index.html")


# --- 手动爬取新闻 ---
@app.route("/api/crawl", methods=["POST"])
def crawl_news():
    """手动触发一次新闻爬取（内置逻辑，无需外部脚本）"""
    from urllib.parse import urljoin
    from bs4 import BeautifulSoup
    import threading

    def do_manual_crawl():
        today = datetime.now().strftime("%Y-%m-%d")
        print(f"[CRAWL] 开始手动爬取 {today}")

        headers_req = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'text/html,application/xhtml+xml',
            'Accept-Language': 'zh-CN,zh;q=0.9',
        }
        sources = [
            ("求是网", "http://www.qstheory.cn/"),
            ("人民网", "http://politics.people.com.cn/"),
            ("新华网", "http://www.xinhuanet.com/politics/"),
        ]
        all_news = {}
        for name, base_url in sources:
            try:
                resp = requests.get(base_url, timeout=15, headers=headers_req, verify=False)
                if resp.status_code == 200:
                    resp.encoding = 'utf-8'
                    soup = BeautifulSoup(resp.text, 'html.parser')
                    count = 0
                    for a in soup.find_all('a', href=True):
                        text = a.get_text(strip=True)
                        href = a['href']
                        if 15 < len(text) < 100 and re.search(r'[一-鿿]', text):
                            exclude = ['登录', '注册', '搜索', '首页', '上一页', '下一页', '评论']
                            if not any(kw in text for kw in exclude):
                                if href.startswith('//'):
                                    href = 'http:' + href
                                elif href.startswith('/'):
                                    href = urljoin(base_url, href)
                                elif not href.startswith('http'):
                                    href = urljoin(base_url, href)
                                if text not in all_news and count < 15:
                                    all_news[text] = {"source": name, "url": href}
                                    count += 1
                print(f"[CRAWL] {name}: {count} 条")
            except Exception as e:
                print(f"[CRAWL] {name} 失败: {e}")

        if not all_news:
            print("[CRAWL] 未获取到新闻")
            return

        priority = {"求是网": 1, "人民网": 2, "新华网": 3}
        sorted_news = sorted(all_news.items(), key=lambda x: priority.get(x[1]["source"], 99))[:15]

        # 保存新闻（AI分析）
        saved_files = []
        for i, (title, info) in enumerate(sorted_news[:8]):
            safe_title = re.sub(r'[\\/*?:"<>|]', '', title)[:40]
            filename = f"{today}_{i+1:02d}_{safe_title}.md"
            filepath = os.path.join(NEWS_ARCHIVE_DIR, filename)
            if os.path.exists(filepath):
                continue
            prompt = f"请根据以下新闻标题，生成一份完整的时政新闻扩展分析（300-500字）：\n\n新闻标题：{title}\n\n请按以下格式输出：\n\n## 📌 新闻背景\n[2-3句背景介绍]\n\n## 📖 核心内容\n[新闻的具体内容扩展]\n\n## 🎯 考公考点\n- 申论角度：[可用主题]\n- 行测考点：[可能出题方向]\n\n## 💡 规范表述\n[2-3句官方标准表述]"
            analysis = call_deepseek(prompt, "你是公考时政分析专家")
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(f"# {title}\n\n")
                f.write(f"**日期**：{today}\n")
                f.write(f"**来源**：{info['source']}\n")
                if info['url']:
                    f.write(f"**原文链接**：{info['url']}\n")
                f.write("\n---\n\n")
                f.write(analysis)
            saved_files.append(filepath)
            print(f"[CRAWL] 已保存: {filename}")

        # 生成日报
        formatted = "\n".join([f"{i+1}. 【{info['source']}】{title}" for i, (title, info) in enumerate(sorted_news)])
        report_prompt = f"""请分析以下时政新闻，按指定格式输出。

【今日新闻】
{formatted}

请按以下格式输出：

📌 **今日时政核心要点**（3-5条）
- [新闻事件]：一句话概括+考公切入点

📝 **一句话考点总结**
1. 【来源】标题 → 考点方向：[具体考点]

🔑 **高频关键词**
关键词1 · 关键词2 · 关键词3

📖 **申论素材卡**
- 可用主题：[主题]
- 核心案例：[新闻事件]
- 规范表述：「官方表述」

🎯 **行测时政预测**
- 可能出题方向：[方向]

💪 **【今日随堂测验】**

请生成2道与今日时政相关的单选题：

**题目1**
[问题]
A. [选项1]  B. [选项2]  C. [选项3]  D. [选项4]
答案：[_]

**题目2**
[问题]
A. [选项1]  B. [选项2]  C. [选项3]  D. [选项4]
答案：[_]
"""
        report = call_deepseek(report_prompt, "你是一个公考时政分析专家。请严格按照要求的格式输出分析结果。")
        report_path = os.path.join(DAILY_REPORTS_DIR, f"daily_news_{today}.md")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(f"# 📰 {today} 时政日报\n\n")
            f.write(report)
        print(f"[CRAWL] 日报已生成: {report_path}")

        # 嵌入知识库
        if saved_files:
            try:
                client = get_chroma_client()
                collection = client.get_or_create_collection(
                    name="shizheng_news",
                    metadata={"description": "考公时政新闻知识库"}
                )
                for filepath in saved_files:
                    with open(filepath, "r", encoding="utf-8") as f:
                        content = f.read()
                    chunks = chunk_text(content)
                    if not chunks:
                        continue
                    embeddings = get_embedding(chunks)
                    if not embeddings:
                        continue
                    fname = os.path.basename(filepath).replace(".md", "")
                    ids = [f"{fname}_chunk_{j}" for j in range(len(chunks))]
                    collection.add(documents=chunks, embeddings=embeddings, ids=ids)
                print(f"[CRAWL] 知识库已更新，新增{len(saved_files)}篇")
            except Exception as e:
                print(f"[CRAWL] 知识库嵌入失败: {e}")

        # 自动同步到GitHub（数据持久化）
        try:
            git_sync_to_github(f"auto: 手动爬取 {today}")
        except Exception as e:
            print(f"[CRAWL] GitHub同步失败: {e}")

        print(f"[CRAWL] 手动爬取完成")

    thread = threading.Thread(target=do_manual_crawl)
    thread.start()
    return jsonify({"success": True, "output": "爬取任务已启动，预计1-2分钟完成。刷新页面查看结果。"})


# --- 新闻归档 ---
@app.route("/api/archives")
def list_archives():
    """按日期分组列出新闻归档"""
    files = glob.glob(os.path.join(NEWS_ARCHIVE_DIR, "*.md"))
    grouped = {}
    for f in sorted(files, reverse=True):
        name = os.path.basename(f)
        match = re.match(r'(\d{4}-\d{2}-\d{2})_(\d+)_(.+)\.md', name)
        if match:
            date, idx, title = match.groups()
            if date not in grouped:
                grouped[date] = []
            grouped[date].append({"index": int(idx), "title": title, "filename": name})
    result = []
    for date in sorted(grouped.keys(), reverse=True):
        items = sorted(grouped[date], key=lambda x: x["index"])
        result.append({"date": date, "count": len(items), "items": items})
    return jsonify(result)


@app.route("/api/archives/<path:filename>")
def get_archive(filename):
    """获取单篇新闻全文，拆分为原文和解析两部分"""
    filepath = os.path.join(NEWS_ARCHIVE_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "文件不存在"}), 404
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    # 提取原文链接
    url_match = re.search(r'\*\*原文链接\*\*：(.+)', content)
    source_url = url_match.group(1).strip() if url_match else ""

    # 将原文链接转为可点击的超链接
    if source_url:
        content = content.replace(f"**原文链接**：{source_url}", f"**原文链接**：[{source_url}]({source_url})")

    # 拆分为原文部分和解析部分
    original_parts = []
    analysis_parts = []
    current_section = "original"
    for line in content.split("\n"):
        if re.match(r'^##\s*🎯\s*考公考点', line) or re.match(r'^##\s*💡\s*规范表述', line):
            current_section = "analysis"
        if current_section == "original":
            original_parts.append(line)
        else:
            analysis_parts.append(line)

    original_md = "\n".join(original_parts).strip()
    analysis_md = "\n".join(analysis_parts).strip()

    html = markdown.markdown(content, extensions=["tables", "fenced_code"])
    # 让所有链接在新标签页打开
    html = html.replace('<a href=', '<a target="_blank" href=')
    original_html = markdown.markdown(original_md, extensions=["tables", "fenced_code"])
    original_html = original_html.replace('<a href=', '<a target="_blank" href=')
    analysis_html = markdown.markdown(analysis_md, extensions=["tables", "fenced_code"])

    return jsonify({
        "filename": filename,
        "markdown": content,
        "html": html,
        "original_html": original_html,
        "analysis_html": analysis_html,
        "source_url": source_url
    })


# --- 日报 ---
@app.route("/api/reports")
def list_reports():
    """列出所有日报和周报，按月分组"""
    files = glob.glob(os.path.join(DAILY_REPORTS_DIR, "*.md"))
    reports = []
    for f in sorted(files, reverse=True):
        name = os.path.basename(f)
        match = re.search(r'(\d{4}-\d{2}-\d{2})', name)
        if match:
            date_str = match.group(1)
            rtype = "周报" if "weekly" in name else "日报"
            reports.append({
                "date": date_str,
                "year": int(date_str[:4]),
                "month": int(date_str[5:7]),
                "day": int(date_str[8:10]),
                "filename": name,
                "type": rtype
            })
    return jsonify(reports)


@app.route("/api/reports/<path:filename>")
def get_report(filename):
    """获取日报内容"""
    filepath = os.path.join(DAILY_REPORTS_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "文件不存在"}), 404
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    html = markdown.markdown(content, extensions=["tables", "fenced_code"])
    return jsonify({"filename": filename, "markdown": content, "html": html})


# --- 知识库查询 ---
@app.route("/api/query", methods=["POST"])
def query_kb():
    """语义查询知识库"""
    data = request.json
    question = data.get("question", "").strip()
    collection_name = data.get("collection", "shizheng_news")
    top_k = data.get("top_k", 5)

    if not question:
        return jsonify({"error": "请输入查询问题"}), 400

    client = get_chroma_client()
    try:
        collection = client.get_collection(collection_name)
    except Exception:
        return jsonify({"error": f"知识库 '{collection_name}' 不存在"}), 404

    embeddings = get_embedding([question])
    if not embeddings:
        return jsonify({"error": "嵌入API调用失败"}), 500

    results = collection.query(query_embeddings=embeddings, n_results=top_k)
    items = []
    for doc, meta, dist in zip(
        results["documents"][0],
        results["metadatas"][0],
        results["distances"][0]
    ):
        items.append({
            "content": doc,
            "title": meta.get("title", ""),
            "date": meta.get("date", ""),
            "source": meta.get("source", ""),
            "score": round(1 - dist, 4)
        })
    return jsonify({"question": question, "results": items})


# --- PDF 做题模块 ---
@app.route("/api/pdf/list")
def list_pdfs():
    """列出已上传的PDF文件"""
    files = glob.glob(os.path.join(PDF_UPLOADS_DIR, "*.pdf"))
    result = []
    for f in sorted(files):
        name = os.path.basename(f)
        size_mb = round(os.path.getsize(f) / 1024 / 1024, 1)
        result.append({"filename": name, "size_mb": size_mb})
    return jsonify(result)


@app.route("/api/pdf/upload", methods=["POST"])
def upload_pdf():
    """上传PDF文件并自动导入知识库"""
    if "file" not in request.files:
        return jsonify({"error": "未选择文件"}), 400
    file = request.files["file"]
    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "仅支持PDF文件"}), 400
    filename = secure_filename(file.filename) or f"upload_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
    # 保留中文文件名
    if file.filename:
        filename = file.filename.replace("/", "_").replace("\\", "_")
    filepath = os.path.join(PDF_UPLOADS_DIR, filename)
    file.save(filepath)

    # 自动导入知识库
    import_result = import_pdf_to_kb(filepath, filename)
    return jsonify({"message": f"上传成功，{import_result}", "filename": filename})


@app.route("/api/pdf/import", methods=["POST"])
def import_pdf():
    """手动将已有PDF导入知识库"""
    data = request.json
    filename = data.get("filename", "")
    filepath = os.path.join(PDF_UPLOADS_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "文件不存在"}), 404
    result = import_pdf_to_kb(filepath, filename)
    return jsonify({"message": result})


def import_pdf_to_kb(filepath: str, filename: str) -> str:
    """将PDF全文向量化并存入知识库"""
    try:
        text = extract_pdf_text(filepath)
        if not text.strip():
            return "PDF无可提取文本，未入库"

        # 分块
        chunks = chunk_text(text, chunk_size=500, overlap=100)
        if not chunks:
            return "文本过短，未入库"

        client = get_chroma_client()
        collection = client.get_or_create_collection(
            name="shizheng_news",
            metadata={"description": "考公时政新闻知识库"}
        )

        # 检查是否已导入过（用文件名前缀去重）
        existing = collection.get()
        existing_ids = set(existing["ids"]) if existing["ids"] else set()
        prefix = filename.replace(".pdf", "")
        if any(eid.startswith(prefix) for eid in existing_ids):
            return "该PDF已导入过知识库"

        # 批量嵌入（每批20条）
        all_embeddings = []
        batch_size = 20
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            embs = get_embedding(batch)
            if not embs:
                return f"嵌入API失败，已处理 {i}/{len(chunks)} 块"
            all_embeddings.extend(embs)

        # 写入
        chunk_ids = [f"{prefix}_p{i}" for i in range(len(chunks))]
        metadatas = [{"title": filename, "date": datetime.now().strftime("%Y-%m-%d"), "source": "PDF导入"} for _ in chunks]
        collection.add(ids=chunk_ids, documents=chunks, embeddings=all_embeddings, metadatas=metadatas)
        return f"已将 {len(chunks)} 段内容导入知识库"
    except Exception as e:
        return f"导入失败: {e}"


@app.route("/api/pdf/extract", methods=["POST"])
def extract_pdf_questions():
    """从PDF中提取指定页范围的内容，让AI生成题目"""
    data = request.json
    filename = data.get("filename", "")
    page_start = data.get("page_start", 1)
    page_end = data.get("page_end", None)
    num_questions = data.get("num_questions", 5)

    filepath = os.path.join(PDF_UPLOADS_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "PDF文件不存在"}), 404

    # 提取指定页范围文本
    doc = fitz.open(filepath)
    total_pages = len(doc)
    page_start = max(1, min(page_start, total_pages))
    if page_end is None:
        page_end = total_pages
    page_end = max(page_start, min(page_end, total_pages))

    text_parts = []
    for i in range(page_start - 1, page_end):
        text_parts.append(doc[i].get_text())
    doc.close()

    content = "\n".join(text_parts).strip()
    if not content:
        return jsonify({"error": "该页范围内无可提取文本"}), 400

    # 截取前4000字避免超长
    if len(content) > 4000:
        content = content[:4000]

    prompt = f"""基于以下PDF文档内容，生成{num_questions}道单选题。

【文档内容】
{content}

请严格按以下JSON格式输出（不要输出其他内容）：
```json
[
  {{
    "id": 1,
    "question": "题目内容",
    "options": {{"A": "选项A", "B": "选项B", "C": "选项C", "D": "选项D"}},
    "answer": "正确答案字母",
    "explanation": "简要解析"
  }}
]
```"""

    system = "你是公务员考试出题专家。根据提供的材料出题，题目要有区分度，选项要有迷惑性。只输出JSON，不要其他文字。"
    result = call_deepseek(prompt, system)

    # 解析JSON
    try:
        # 提取JSON部分
        json_match = re.search(r'\[[\s\S]*\]', result)
        if json_match:
            questions = json.loads(json_match.group())
        else:
            questions = json.loads(result)
    except json.JSONDecodeError:
        return jsonify({
            "error": "AI返回格式异常，请重试",
            "raw": result
        }), 500

    return jsonify({
        "filename": filename,
        "pages": f"{page_start}-{page_end}",
        "total_pages": total_pages,
        "questions": questions
    })


@app.route("/api/pdf/check", methods=["POST"])
def check_answer():
    """AI判断用户答案并给出详细解析"""
    data = request.json
    question = data.get("question", "")
    options = data.get("options", {})
    correct_answer = data.get("correct_answer", "")
    user_answer = data.get("user_answer", "")
    explanation = data.get("explanation", "")

    is_correct = user_answer.upper() == correct_answer.upper()

    # 调用AI给出详细解析
    options_text = "\n".join([f"{k}. {v}" for k, v in options.items()])
    prompt = f"""题目：{question}
{options_text}

正确答案：{correct_answer}
用户选择：{user_answer}
用户{'回答正确' if is_correct else '回答错误'}。

请给出详细的解题分析：
1. 逐项分析每个选项的对错原因
2. 解释为什么正确答案是{correct_answer}
3. 给出相关知识点的记忆技巧"""

    system = "你是公务员考试辅导专家，请给出清晰、有条理的解题分析。"
    analysis = call_deepseek(prompt, system)
    analysis_html = markdown.markdown(analysis, extensions=["tables", "fenced_code"])

    return jsonify({
        "is_correct": is_correct,
        "correct_answer": correct_answer,
        "user_answer": user_answer,
        "analysis": analysis,
        "analysis_html": analysis_html
    })


# --- PDF题目提取刷题模块 ---
PDF_QUESTIONS_DIR = os.path.join(DATA_DIR, "pdf_questions")
os.makedirs(PDF_QUESTIONS_DIR, exist_ok=True)


def ai_generate_questions_from_pdf(filepath: str) -> list:
    """当PDF无可提取文本时，使用AI根据PDF内容生成题目
    对于扫描件等图片PDF，先尝试用fitz提取图片描述，再用AI出题
    """
    # 尝试提取PDF中每页的文本（即使很少）
    doc = fitz.open(filepath)
    page_count = len(doc)
    # 获取文件名作为主题参考
    filename = os.path.basename(filepath)
    topic = re.sub(r'\.(pdf|PDF)$', '', filename)
    topic = re.sub(r'[_\-]', ' ', topic)

    # 尝试提取少量文本作为上下文
    sparse_text = ""
    for page in doc:
        t = page.get_text().strip()
        if t:
            sparse_text += t + "\n"
    doc.close()

    # 构建prompt
    if sparse_text and len(sparse_text) > 50:
        # 有部分文本但不足以直接提取题目
        context_hint = f"以下是从PDF中提取到的部分文本内容：\n{sparse_text[:3000]}\n\n"
    else:
        # 完全无文本（纯图片PDF）
        context_hint = f"这是一份名为《{topic}》的PDF文档（共{page_count}页），无法提取文本内容。\n\n"

    prompt = f"""{context_hint}请根据以上信息，生成20道公务员考试时政类选择题。

要求：
1. 每道题必须有4个选项（A、B、C、D）
2. 必须标注正确答案
3. 题目应涵盖时政热点、政策法规、重要会议等内容
4. 题目难度适中，符合公务员考试水平

请严格按照以下JSON格式输出（不要输出其他内容）：
[
  {{
    "id": 1,
    "question": "题干内容",
    "options": {{"A": "选项A", "B": "选项B", "C": "选项C", "D": "选项D"}},
    "answer": "正确答案字母"
  }},
  ...
]"""

    system = "你是公务员考试命题专家，擅长根据时政材料出选择题。请只输出JSON数组，不要输出任何其他文字。"
    result = call_deepseek(prompt, system)

    # 解析AI返回的JSON
    questions = _parse_ai_questions_response(result)
    return questions


def _parse_ai_questions_response(text: str) -> list:
    """解析AI返回的题目JSON"""
    # 尝试直接解析
    text = text.strip()
    # 去除可能的markdown代码块标记
    if text.startswith("```"):
        text = re.sub(r'^```(?:json)?\s*\n?', '', text)
        text = re.sub(r'\n?```\s*$', '', text)
    text = text.strip()

    try:
        questions = json.loads(text)
        if isinstance(questions, list):
            # 验证格式
            valid = []
            for q in questions:
                if all(k in q for k in ("id", "question", "options", "answer")):
                    if isinstance(q["options"], dict) and len(q["options"]) >= 2:
                        valid.append(q)
            return valid
    except json.JSONDecodeError:
        pass

    # 尝试从文本中提取JSON数组
    match = re.search(r'\[[\s\S]*\]', text)
    if match:
        try:
            questions = json.loads(match.group())
            if isinstance(questions, list):
                valid = []
                for q in questions:
                    if all(k in q for k in ("id", "question", "options", "answer")):
                        if isinstance(q["options"], dict) and len(q["options"]) >= 2:
                            valid.append(q)
                return valid
        except json.JSONDecodeError:
            pass

    return []


def parse_questions_from_text(text: str) -> list:
    """从PDF文本中提取完整的选择题，保证题目完整性
    支持两种格式：
    1. 题号在题干前面：1. 题干文本\nA.选项...
    2. 题号在题干后面（粉笔等）：题干文本\n1.\nA.选项...
    """
    questions = []
    lines = text.split('\n')

    # 先尝试检测格式：题号是否单独一行
    # 找所有 "数字." 单独成行的位置
    num_line_indices = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if re.match(r'^\d{1,3}\.\s*$', stripped):
            num_line_indices.append(i)

    if len(num_line_indices) >= 5:
        # 格式2：题号单独一行（粉笔格式）
        questions = _parse_format_fenbi(lines, num_line_indices)
    else:
        # 格式1：题号在题干开头
        questions = _parse_format_standard(text)

    return questions


def _parse_format_fenbi(lines: list, num_indices: list) -> list:
    """解析粉笔格式：题干在题号前面，选项在题号后面"""
    questions = []
    # 过滤掉页眉页脚行
    skip_patterns = [r'^·\s*本试卷', r'^第\s*\d+\s*页', r'^\d{4}年.*模考', r'^\d{4}年.*大赛', r'^.*执法类\）?\s*$']

    for idx_pos, line_idx in enumerate(num_indices):
        num_match = re.match(r'^(\d{1,3})\.\s*$', lines[line_idx].strip())
        if not num_match:
            continue
        num = int(num_match.group(1))

        # 提取选项（题号行之后）
        options = {}
        opt_end = line_idx + 1
        for j in range(line_idx + 1, min(line_idx + 20, len(lines))):
            opt_match = re.match(r'^([A-D])[.．]\s*(.+)', lines[j].strip())
            if opt_match:
                options[opt_match.group(1)] = opt_match.group(2).strip()
                opt_end = j + 1
            elif options and not re.match(r'^[A-D][.．]', lines[j].strip()):
                # 选项可能换行续接
                if lines[j].strip() and not re.match(r'^\d{1,3}\.\s*$', lines[j].strip()):
                    last_key = list(options.keys())[-1]
                    options[last_key] += lines[j].strip()
                else:
                    break

        if len(options) < 2:
            continue

        # 提取题干（从上一个题目的选项结束到当前题号之间）
        if idx_pos == 0:
            stem_start = 0
        else:
            # 从上一题的选项结束位置开始
            prev_num_idx = num_indices[idx_pos - 1]
            stem_start = prev_num_idx + 1
            # 跳过上一题的选项行
            for j in range(prev_num_idx + 1, line_idx):
                if re.match(r'^[A-D][.．]', lines[j].strip()):
                    stem_start = j + 1
                elif lines[j].strip() and stem_start <= j and not re.match(r'^[A-D][.．]', lines[j].strip()):
                    break

        # 收集题干行
        stem_lines = []
        for j in range(stem_start, line_idx):
            l = lines[j].strip()
            if not l:
                continue
            # 跳过页眉页脚
            if any(re.match(p, l) for p in skip_patterns):
                continue
            # 跳过章节标题（如 "一. 政治理论：..."）
            if re.match(r'^[一二三四五六七八九十]+[.．、]', l):
                continue
            # 跳过上一题的选项残留
            if re.match(r'^[A-D][.．]', l):
                continue
            stem_lines.append(l)

        stem = ' '.join(stem_lines).strip()
        # 清理多余空格
        stem = re.sub(r'\s+', ' ', stem)

        if len(stem) < 5:
            continue

        # 检查答案
        answer = ""
        for j in range(opt_end, min(opt_end + 3, len(lines))):
            ans_match = re.search(r'答案[：:]\s*\[?([A-D])\]?', lines[j] if j < len(lines) else "")
            if ans_match:
                answer = ans_match.group(1)
                break

        questions.append({
            "id": num,
            "question": stem,
            "options": options,
            "answer": answer
        })

    return questions


def _parse_format_standard(text: str) -> list:
    """解析标准格式：题号在题干开头"""
    questions = []
    parts = re.split(r'\n(?=\d{1,3}[.．]\s*[^\s\d])', text)

    for part in parts:
        part = part.strip()
        if not part:
            continue
        num_match = re.match(r'^(\d{1,3})[.．]\s*([\s\S]*)', part)
        if not num_match:
            continue

        num = int(num_match.group(1))
        body = num_match.group(2).strip()
        if not body:
            continue

        # 提取选项
        option_matches = re.findall(r'(?:^|\n)\s*([A-D])[.．、]\s*(.+?)(?=(?:\n\s*[A-D][.．、])|(?:\n\s*答案)|$)', body, re.DOTALL)
        if len(option_matches) < 2:
            continue

        # 题干
        first_opt = re.search(r'(?:^|\n)\s*A[.．、]\s*', body)
        if first_opt:
            stem = body[:first_opt.start()].strip()
        else:
            continue

        stem = re.sub(r'\s+', ' ', stem).strip()
        if len(stem) < 5:
            continue

        options = {}
        for key, val in option_matches:
            options[key] = re.sub(r'\s+', ' ', val).strip()

        if 'A' not in options or 'B' not in options:
            continue

        answer = ""
        ans_match = re.search(r'答案[：:]\s*\[?([A-D])\]?', body)
        if ans_match:
            answer = ans_match.group(1)

        questions.append({
            "id": num,
            "question": stem,
            "options": options,
            "answer": answer
        })

    return questions


@app.route("/api/pdf/extract-questions", methods=["POST"])
def extract_questions_from_pdf():
    """从PDF中提取所有选择题，保存到本地JSON"""
    data = request.json
    filename = data.get("filename", "")
    filepath = os.path.join(PDF_UPLOADS_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "PDF文件不存在"}), 404

    # 检查是否已提取过
    safe_name = re.sub(r'[^\w\-.]', '_', filename)
    json_path = os.path.join(PDF_QUESTIONS_DIR, safe_name + ".json")
    if os.path.exists(json_path):
        with open(json_path, "r", encoding="utf-8") as f:
            existing = json.load(f)
        return jsonify({"message": f"已提取过，共{len(existing)}题", "count": len(existing), "questions": existing})

    # 提取PDF全文
    text = extract_pdf_text(filepath)
    ai_generated = False

    if not text.strip():
        # 无可提取文本，使用AI生成题目
        questions = ai_generate_questions_from_pdf(filepath)
        ai_generated = True
    else:
        # 解析题目
        questions = parse_questions_from_text(text)
        if not questions:
            # 有文本但无法识别题目格式，也用AI出题
            questions = ai_generate_questions_from_pdf(filepath)
            ai_generated = True

    if not questions:
        return jsonify({"error": "未能提取或生成题目，请检查PDF内容"}), 400

    # 保存
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)

    # 同步到GitHub
    try:
        import threading as _t
        _t.Thread(target=git_sync_to_github, args=(f"auto: 提取题目 {filename}",), daemon=True).start()
    except Exception:
        pass

    msg = f"AI生成{len(questions)}道题目" if ai_generated else f"成功提取{len(questions)}道题目"
    return jsonify({"message": msg, "count": len(questions), "questions": questions, "ai_generated": ai_generated})


@app.route("/api/pdf/questions")
def get_pdf_questions():
    """获取某PDF已提取的题目"""
    filename = request.args.get("filename", "")
    safe_name = re.sub(r'[^\w\-.]', '_', filename)
    json_path = os.path.join(PDF_QUESTIONS_DIR, safe_name + ".json")
    if not os.path.exists(json_path):
        return jsonify({"questions": [], "count": 0})
    with open(json_path, "r", encoding="utf-8") as f:
        questions = json.load(f)
    return jsonify({"questions": questions, "count": len(questions)})


@app.route("/api/pdf/questions/analyze", methods=["POST"])
def analyze_question():
    """AI解析单道题目"""
    data = request.json
    question = data.get("question", "")
    options = data.get("options", {})
    user_answer = data.get("user_answer", "")
    correct_answer = data.get("correct_answer", "")

    options_text = "\n".join([f"{k}. {v}" for k, v in options.items()])

    if correct_answer:
        prompt = f"""题目：{question}
{options_text}

正确答案：{correct_answer}
用户选择：{user_answer}

请给出详细解析：
1. 逐项分析每个选项
2. 解释正确答案的原因
3. 给出相关知识点和记忆技巧"""
    else:
        prompt = f"""题目：{question}
{options_text}

用户选择：{user_answer}

请分析这道题：
1. 判断正确答案应该是哪个
2. 逐项分析每个选项的对错
3. 给出解题思路和相关知识点"""

    system = "你是公务员考试辅导专家，请给出清晰、有条理的解题分析。"
    analysis = call_deepseek(prompt, system)
    analysis_html = markdown.markdown(analysis, extensions=["tables", "fenced_code"])

    # 如果没有标准答案，尝试从AI回复中提取
    ai_answer = ""
    if not correct_answer:
        ans_match = re.search(r'正确答案[是为：:]+\s*([A-D])', analysis)
        if ans_match:
            ai_answer = ans_match.group(1)

    return jsonify({
        "analysis": analysis,
        "analysis_html": analysis_html,
        "ai_answer": ai_answer
    })


@app.route("/api/pdf/questions/reextract", methods=["POST"])
def reextract_questions():
    """重新提取题目（删除旧缓存）"""
    data = request.json
    filename = data.get("filename", "")
    safe_name = re.sub(r'[^\w\-.]', '_', filename)
    json_path = os.path.join(PDF_QUESTIONS_DIR, safe_name + ".json")
    if os.path.exists(json_path):
        os.remove(json_path)
    # 重新提取
    filepath = os.path.join(PDF_UPLOADS_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "PDF文件不存在"}), 404
    text = extract_pdf_text(filepath)
    ai_generated = False

    if not text.strip():
        questions = ai_generate_questions_from_pdf(filepath)
        ai_generated = True
    else:
        questions = parse_questions_from_text(text)
        if not questions:
            questions = ai_generate_questions_from_pdf(filepath)
            ai_generated = True

    if not questions:
        return jsonify({"error": "未能提取或生成题目"}), 400
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)
    msg = f"AI重新生成{len(questions)}道题" if ai_generated else f"重新提取成功，共{len(questions)}道题"
    return jsonify({"message": msg, "count": len(questions), "questions": questions, "ai_generated": ai_generated})


# --- 知识库管理 ---
@app.route("/api/collections")
def list_collections():
    """列出所有知识库集合"""
    client = get_chroma_client()
    collections = client.list_collections()
    result = []
    for col in collections:
        c = client.get_collection(col.name)
        meta = col.metadata or {}
        display = meta.get("display_name", None)
        if not display:
            # 默认知识库显示中文名
            if col.name == "shizheng_news":
                display = "时政新闻库"
            else:
                display = col.name
        result.append({
            "name": col.name,
            "display_name": display,
            "count": c.count(),
            "metadata": meta
        })
    return jsonify(result)


@app.route("/api/collections", methods=["POST"])
def create_collection():
    """创建新的知识库集合"""
    data = request.json
    display_name = data.get("name", "").strip()
    description = data.get("description", "")
    if not display_name:
        return jsonify({"error": "请输入知识库名称"}), 400

    # 转换为合法的集合名
    col_name = sanitize_collection_name(display_name)

    client = get_chroma_client()
    try:
        collection = client.get_or_create_collection(
            name=col_name, metadata={"description": description, "display_name": display_name}
        )
        return jsonify({"name": col_name, "display_name": display_name, "count": collection.count(), "message": "创建成功"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/collections/<name>/upload", methods=["POST"])
def upload_to_collection(name):
    """上传文本到指定知识库"""
    data = request.json
    text = data.get("text", "").strip()
    title = data.get("title", "手动添加")
    if not text:
        return jsonify({"error": "请输入文本内容"}), 400

    client = get_chroma_client()
    try:
        collection = client.get_collection(name)
    except Exception:
        return jsonify({"error": f"知识库 '{name}' 不存在"}), 404

    chunks = chunk_text(text)
    chunk_ids = [f"{title}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{i}" for i in range(len(chunks))]
    metadatas = [{"title": title, "date": datetime.now().strftime("%Y-%m-%d"), "source": "手动上传"} for _ in chunks]

    embeddings = get_embedding(chunks)
    if not embeddings:
        return jsonify({"error": "嵌入API调用失败"}), 500

    collection.add(ids=chunk_ids, documents=chunks, embeddings=embeddings, metadatas=metadatas)
    return jsonify({"message": f"已添加 {len(chunks)} 个文本块", "total": collection.count()})


@app.route("/api/collections/<name>", methods=["DELETE"])
def delete_collection(name):
    """删除知识库集合"""
    if name == "shizheng_news":
        return jsonify({"error": "默认知识库不可删除"}), 400
    client = get_chroma_client()
    try:
        client.delete_collection(name)
        return jsonify({"message": f"已删除知识库 '{name}'"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- 错题本 ---
@app.route("/api/wrong-book")
def get_wrong_book():
    """获取错题列表"""
    data = load_user_data()
    return jsonify(data.get("wrong_questions", []))


@app.route("/api/wrong-book", methods=["POST"])
def add_wrong_question():
    """添加错题"""
    data = load_user_data()
    item = request.json
    item["id"] = hashlib.md5(f"{item.get('question','')}{datetime.now()}".encode()).hexdigest()[:12]
    item["time"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    data.setdefault("wrong_questions", []).insert(0, item)
    save_user_data(data)
    return jsonify({"message": "已加入错题本"})


@app.route("/api/wrong-book/<qid>", methods=["DELETE"])
def remove_wrong_question(qid):
    """移除错题（已掌握）"""
    data = load_user_data()
    data["wrong_questions"] = [q for q in data.get("wrong_questions", []) if q.get("id") != qid]
    save_user_data(data)
    return jsonify({"message": "已移除"})


# --- 做题记录 ---
@app.route("/api/quiz-history", methods=["POST"])
def add_quiz_history():
    """记录一次做题"""
    data = load_user_data()
    record = request.json
    record["time"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    data.setdefault("quiz_history", []).insert(0, record)
    # 只保留最近200条
    data["quiz_history"] = data["quiz_history"][:200]
    save_user_data(data)
    return jsonify({"message": "已记录"})


# --- 收藏 ---
@app.route("/api/favorites")
def get_favorites():
    """获取收藏列表"""
    data = load_user_data()
    return jsonify(data.get("favorites", []))


@app.route("/api/favorites", methods=["POST"])
def add_favorite():
    """收藏文章"""
    data = load_user_data()
    item = request.json
    # 去重
    existing = [f["filename"] for f in data.get("favorites", [])]
    if item.get("filename") in existing:
        return jsonify({"message": "已收藏过"})
    item["time"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    data.setdefault("favorites", []).insert(0, item)
    save_user_data(data)
    return jsonify({"message": "收藏成功"})


@app.route("/api/favorites/<path:filename>", methods=["DELETE"])
def remove_favorite(filename):
    """取消收藏"""
    data = load_user_data()
    data["favorites"] = [f for f in data.get("favorites", []) if f.get("filename") != filename]
    save_user_data(data)
    return jsonify({"message": "已取消收藏"})


# --- 学习统计 ---
@app.route("/api/stats")
def get_stats():
    """获取学习统计"""
    data = load_user_data()
    history = data.get("quiz_history", [])
    wrong = data.get("wrong_questions", [])

    total_quizzes = len(history)
    total_questions = sum(h.get("total", 0) for h in history)
    total_correct = sum(h.get("correct", 0) for h in history)
    accuracy = round(total_correct / total_questions * 100, 1) if total_questions else 0

    # 近7天每日统计
    daily = {}
    for h in history:
        day = h.get("time", "")[:10]
        if day:
            if day not in daily:
                daily[day] = {"total": 0, "correct": 0}
            daily[day]["total"] += h.get("total", 0)
            daily[day]["correct"] += h.get("correct", 0)

    # 取最近7天
    from datetime import timedelta
    today = datetime.now().date()
    week_data = []
    for i in range(6, -1, -1):
        d = (today - timedelta(days=i)).strftime("%Y-%m-%d")
        info = daily.get(d, {"total": 0, "correct": 0})
        week_data.append({"date": d, "total": info["total"], "correct": info["correct"]})

    return jsonify({
        "total_quizzes": total_quizzes,
        "total_questions": total_questions,
        "total_correct": total_correct,
        "accuracy": accuracy,
        "wrong_count": len(wrong),
        "week_data": week_data
    })


# --- 知识查询追问 ---
@app.route("/api/daily-quiz")
def get_daily_quiz():
    """从最新日报中提取今日随堂测验，返回结构化题目数据"""
    from datetime import timedelta
    for offset in range(3):
        d = (datetime.now() - timedelta(days=offset)).strftime("%Y-%m-%d")
        filepath = os.path.join(DAILY_REPORTS_DIR, f"daily_news_{d}.md")
        if os.path.exists(filepath):
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
            # 提取随堂测验部分
            quiz_match = re.search(r'今日随堂测验.*?\n(.*?)(?=━━━|📌\s*\*\*明日预告|$)', content, re.DOTALL)
            if not quiz_match:
                continue

            quiz_text = quiz_match.group(1).strip()
            # 解析题目
            questions = []
            q_blocks = re.split(r'\*\*题目\d+\*\*\s*\n?', quiz_text)
            for block in q_blocks:
                block = block.strip()
                if not block:
                    continue
                # 提取题干
                lines = block.split('\n')
                q_text = lines[0].strip() if lines else ''
                # 提取选项
                options = {}
                for line in lines[1:]:
                    opt_match = re.match(r'^([A-D])[.、]\s*(.+)', line.strip())
                    if opt_match:
                        options[opt_match.group(1)] = opt_match.group(2).strip()
                # 提取答案
                ans_match = re.search(r'答案[：:]\s*\[?([A-D])\]?', block)
                answer = ans_match.group(1) if ans_match else ''

                if q_text and options and answer:
                    questions.append({
                        "id": len(questions) + 1,
                        "question": q_text,
                        "options": options,
                        "answer": answer
                    })

            if questions:
                return jsonify({"date": d, "questions": questions})
    return jsonify({"date": "", "questions": []})


@app.route("/api/query/explain", methods=["POST"])
def query_explain():
    """对查询结果进行AI追问解读"""
    data = request.json
    question = data.get("question", "")
    context = data.get("context", "")

    prompt = f"""用户查询了"{question}"，以下是知识库中检索到的相关内容：

{context}

请基于以上内容，为用户做一个清晰的总结和深入解读：
1. 概括核心要点
2. 分析考公相关考点
3. 给出记忆建议"""

    system = "你是公务员考试辅导专家，请基于提供的材料给出精准、有条理的解读。"
    answer = call_deepseek(prompt, system)
    html = markdown.markdown(answer, extensions=["tables", "fenced_code"])
    return jsonify({"answer": answer, "html": html})


# ========== 申论训练模块 ==========

def _search_kb_for_essay(text: str, top_k: int = 5) -> str:
    """从知识库检索与文本相关的素材"""
    client = get_chroma_client()
    try:
        collection = client.get_collection("shizheng_news")
    except Exception:
        return ""

    # 提取关键句作为查询
    query_text = text[:200] if len(text) > 200 else text
    embeddings = get_embedding([query_text])
    if not embeddings:
        return ""

    results = collection.query(query_embeddings=embeddings, n_results=top_k)
    if not results["documents"][0]:
        return ""

    materials = []
    for i, doc in enumerate(results["documents"][0]):
        materials.append(f"【素材{i+1}】{doc[:300]}")
    return "\n\n".join(materials)


@app.route("/api/essay/review", methods=["POST"])
def essay_review():
    """申论批改打分"""
    data = request.json
    essay = data.get("essay", "").strip()
    topic = data.get("topic", "").strip()

    if not essay:
        return jsonify({"error": "请输入申论内容"}), 400
    if len(essay) < 50:
        return jsonify({"error": "文章太短，至少50字"}), 400

    # 从知识库检索相关素材
    kb_materials = _search_kb_for_essay(essay)

    topic_line = f"\n申论题目：{topic}" if topic else ""
    materials_line = f"\n\n【知识库参考素材】\n{kb_materials}" if kb_materials else ""

    prompt = f"""请批改以下申论文章，按照公务员考试评分标准进行评价。{topic_line}

【用户文章】
{essay}
{materials_line}

请严格按以下格式输出：

## 📊 总体评分

| 维度 | 分数 | 简评 |
|------|------|------|
| 结构布局 | /20 | |
| 论点深度 | /25 | |
| 素材运用 | /20 | |
| 语言表达 | /20 | |
| 政策理论 | /15 | |
| **总分** | **/100** | |

## 📝 整体评价
[2-3句总结性评价]

## 🔍 逐段批注

[对文章各段落逐一给出具体修改建议，指出优点和问题]

## ✅ 修改建议

[列出3-5条最重要的改进方向]

## 💡 推荐素材

[基于知识库素材，推荐2-3个可以引用的时政案例/表述]"""

    system = "你是资深公考申论阅卷专家，有10年申论批改经验。请按照国考申论评分标准进行客观、专业的批改，指出不足的同时肯定优点。"
    answer = call_deepseek(prompt, system)
    html = markdown.markdown(answer, extensions=["tables", "fenced_code"])
    return jsonify({"answer": answer, "html": html, "has_kb_materials": bool(kb_materials)})


@app.route("/api/essay/assist", methods=["POST"])
def essay_assist():
    """AI辅助写作"""
    data = request.json
    topic = data.get("topic", "").strip()
    material = data.get("material", "").strip()
    word_count = data.get("word_count", 1000)

    if not topic:
        return jsonify({"error": "请输入申论题目"}), 400

    # 从知识库检索相关素材
    kb_materials = _search_kb_for_essay(topic + " " + material)

    material_line = f"\n\n【背景材料】\n{material}" if material else ""
    kb_line = f"\n\n【知识库素材（请融入文章）】\n{kb_materials}" if kb_materials else ""

    prompt = f"""请根据以下申论题目撰写一篇示范申论。

【题目】{topic}{material_line}{kb_line}

要求：
1. 字数约{word_count}字
2. 严格按照申论格式：标题 → 开头（引出主题）→ 分论点（2-3个）→ 结尾（总结升华）
3. 融入时政热点和官方表述
4. 语言规范，逻辑清晰
5. 标注引用的素材来源

请按以下格式输出：

## ✍️ 示范申论

[完整申论文章]

---

## 📌 写作思路解析

- **标题**：[解释标题拟定思路]
- **结构**：[解释分论点设置逻辑]
- **素材运用**：[列出引用的时政案例]

## 🎯 关键表述参考

[列出3-5句可直接使用的官方规范表述]"""

    system = "你是公务员考试申论辅导名师，擅长撰写高分申论范文。请结合最新时政，写出逻辑严谨、语言规范、有深度的申论。"
    answer = call_deepseek(prompt, system)
    html = markdown.markdown(answer, extensions=["tables", "fenced_code"])
    return jsonify({"answer": answer, "html": html, "has_kb_materials": bool(kb_materials)})


@app.route("/api/essay/polish", methods=["POST"])
def essay_polish():
    """申论润色优化"""
    data = request.json
    essay = data.get("essay", "").strip()
    style = data.get("style", "政论")

    if not essay:
        return jsonify({"error": "请输入需要润色的文章"}), 400
    if len(essay) < 50:
        return jsonify({"error": "文章太短，至少50字"}), 400

    # 从知识库检索相关素材
    kb_materials = _search_kb_for_essay(essay)
    kb_line = f"\n\n【知识库参考素材（可用于补充）】\n{kb_materials}" if kb_materials else ""

    prompt = f"""请将以下文章按照「{style}」风格进行润色优化，使其更符合公务员申论写作规范。

【原文】
{essay}{kb_line}

润色要求：
1. 保持原文核心观点不变
2. 优化语言表达，使用官方规范用语
3. 加强论证逻辑，补充论据
4. 如有相关时政素材，适当融入
5. 纠正语病和不规范表述

请按以下格式输出：

## ✨ 润色后文章

[完整的润色后文章]

---

## 📋 修改说明

| 修改位置 | 原文 | 修改后 | 修改原因 |
|----------|------|--------|----------|
| ... | ... | ... | ... |

## 💎 提升要点

[总结3-5条主要提升点]"""

    system = f"你是资深申论写作教练，专长{style}类文体。请在保持原文核心观点的基础上进行专业润色，提升文章质量。"
    answer = call_deepseek(prompt, system)
    html = markdown.markdown(answer, extensions=["tables", "fenced_code"])
    return jsonify({"answer": answer, "html": html, "has_kb_materials": bool(kb_materials)})


@app.route("/api/essay/extract-topic", methods=["POST"])
def essay_extract_topic():
    """上传题目文件(PDF/Word)，提取并整理出申论题目与背景材料"""
    if "file" not in request.files:
        return jsonify({"error": "未选择文件"}), 400
    file = request.files["file"]
    fname = file.filename or ""
    if not fname.lower().endswith((".pdf", ".docx", ".txt", ".md")):
        return jsonify({"error": "仅支持 PDF、Word(.docx) 或文本文件"}), 400

    safe = fname.replace("/", "_").replace("\\", "_")
    filepath = os.path.join(PDF_UPLOADS_DIR, f"essay_topic_{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe}")
    file.save(filepath)

    try:
        raw_text = extract_document_text(filepath, fname)
    except Exception as e:
        return jsonify({"error": f"文件解析失败: {e}"}), 400
    finally:
        try:
            os.remove(filepath)
        except Exception:
            pass

    raw_text = raw_text.strip()
    if not raw_text:
        return jsonify({"error": "未能从文件中提取到文本（可能是扫描件图片）"}), 400

    # 截断，避免超长
    snippet = raw_text[:6000]
    prompt = f"""下面是从用户上传的申论题目文件中提取的原始文本，可能包含背景材料、作答要求等。请你整理出结构化的题目信息。

【原始文本】
{snippet}

请严格输出 JSON（不要任何额外解释、不要markdown代码块包裹）：
{{
  "topic": "申论作答的核心题目/写作要求（如果是大作文题，提炼出写作主题与要求）",
  "material": "题目附带的背景材料摘要（精简到400字以内，保留关键数据与事例）",
  "requirements": "字数、文体等具体作答要求"
}}"""
    system = "你是申论命题解析专家，擅长从题目文件中精准提取写作要求和背景材料。只输出合法JSON。"
    answer = call_deepseek(prompt, system)

    # 解析JSON
    parsed = _parse_json_loose(answer)
    if not parsed:
        # 兜底：把全文当material返回
        return jsonify({
            "topic": "",
            "material": raw_text[:2000],
            "requirements": "",
            "raw_len": len(raw_text),
            "note": "未能自动结构化，已返回原文供你手动确认"
        })
    parsed["raw_len"] = len(raw_text)
    return jsonify(parsed)


@app.route("/api/essay/compare", methods=["POST"])
def essay_compare():
    """根据题目生成范文，并与用户作文逐项对比、给出修改建议"""
    data = request.json
    topic = data.get("topic", "").strip()
    material = data.get("material", "").strip()
    user_essay = data.get("essay", "").strip()
    word_count = data.get("word_count", 1000)

    if not topic and not material:
        return jsonify({"error": "请先提供题目或背景材料"}), 400
    if not user_essay:
        return jsonify({"error": "请输入你的作文"}), 400
    if len(user_essay) < 50:
        return jsonify({"error": "你的作文太短，至少50字"}), 400

    # 从知识库检索相关时政素材
    kb_materials = _search_kb_for_essay((topic + " " + material + " " + user_essay)[:500])

    material_line = f"\n\n【背景材料】\n{material}" if material else ""
    kb_line = f"\n\n【知识库时政素材（写范文时请融入）】\n{kb_materials}" if kb_materials else ""

    prompt = f"""你将完成两项任务：先根据题目撰写一篇高分范文，再把它与用户的作文对比，给出精准的修改方案。

【题目/写作要求】
{topic}{material_line}{kb_line}

【用户的作文】
{user_essay}

请严格按以下 Markdown 格式输出：

## ✍️ 高分范文

[根据题目撰写一篇约{word_count}字的范文，标题居中，结构完整：开头引题 → 分论点(2-3个) → 结尾升华，融入时政素材]

---

## 🔍 对比分析

| 维度 | 你的作文 | 范文 | 差距 |
|------|----------|------|------|
| 立意角度 | | | |
| 结构布局 | | | |
| 论点深度 | | | |
| 素材运用 | | | |
| 语言表达 | | | |

## 📝 逐段修改建议

[针对用户作文的各段落，指出问题并给出具体改写示例，格式为：
**第X段问题**：...
**建议改为**：...]

## 🎯 提分要点

[列出3-5条最关键的提升方向，按重要性排序]

## 💎 可直接套用的金句

[从范文中提炼3-5句可背诵套用的规范表述]"""

    system = "你是国考申论阅卷名师，既能写出高分范文，又能精准诊断学员作文的差距。批改客观专业，修改建议具体可操作。"
    answer = call_deepseek(prompt, system)
    html = markdown.markdown(answer, extensions=["tables", "fenced_code"])
    return jsonify({"answer": answer, "html": html, "has_kb_materials": bool(kb_materials)})


# ========== 定时爬取接口 ==========
CRON_SECRET = os.environ.get("CRON_SECRET", "shizheng2026")

@app.route("/api/cron/daily-crawl", methods=["POST", "GET"])
def cron_daily_crawl():
    """定时爬取新闻，供外部cron服务调用"""
    # 简单鉴权
    secret = request.args.get("secret", "") or request.headers.get("X-Cron-Secret", "")
    if secret != CRON_SECRET:
        return jsonify({"error": "unauthorized"}), 401

    from urllib.parse import urljoin
    from bs4 import BeautifulSoup
    import threading

    def do_crawl():
        today = datetime.now().strftime("%Y-%m-%d")
        print(f"[CRON] 开始每日爬取 {today}")

        # 1. 爬取新闻
        headers_req = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'text/html,application/xhtml+xml',
            'Accept-Language': 'zh-CN,zh;q=0.9',
        }
        sources = [
            ("求是网", "http://www.qstheory.cn/"),
            ("人民网", "http://politics.people.com.cn/"),
            ("新华网", "http://www.xinhuanet.com/politics/"),
        ]
        all_news = {}
        for name, base_url in sources:
            try:
                resp = requests.get(base_url, timeout=15, headers=headers_req, verify=False)
                if resp.status_code == 200:
                    resp.encoding = 'utf-8'
                    soup = BeautifulSoup(resp.text, 'html.parser')
                    count = 0
                    for a in soup.find_all('a', href=True):
                        text = a.get_text(strip=True)
                        href = a['href']
                        if 15 < len(text) < 100 and re.search(r'[一-鿿]', text):
                            exclude = ['登录', '注册', '搜索', '首页', '上一页', '下一页', '评论']
                            if not any(kw in text for kw in exclude):
                                if href.startswith('//'):
                                    href = 'http:' + href
                                elif href.startswith('/'):
                                    href = urljoin(base_url, href)
                                elif not href.startswith('http'):
                                    href = urljoin(base_url, href)
                                if text not in all_news and count < 15:
                                    all_news[text] = {"source": name, "url": href}
                                    count += 1
                print(f"[CRON] {name}: {count} 条")
            except Exception as e:
                print(f"[CRON] {name} 失败: {e}")

        if not all_news:
            print("[CRON] 未获取到新闻")
            return

        # 排序
        priority = {"求是网": 1, "人民网": 2, "新华网": 3}
        sorted_news = sorted(all_news.items(), key=lambda x: priority.get(x[1]["source"], 99))[:15]

        # 2. 保存新闻全文（AI分析）
        saved_files = []
        for i, (title, info) in enumerate(sorted_news[:8]):
            safe_title = re.sub(r'[\\/*?:"<>|]', '', title)[:40]
            filename = f"{today}_{i+1:02d}_{safe_title}.md"
            filepath = os.path.join(NEWS_ARCHIVE_DIR, filename)
            if os.path.exists(filepath):
                continue
            # AI生成分析
            prompt = f"请根据以下新闻标题，生成一份完整的时政新闻扩展分析（300-500字）：\n\n新闻标题：{title}\n\n请按以下格式输出：\n\n## 📌 新闻背景\n[2-3句背景介绍]\n\n## 📖 核心内容\n[新闻的具体内容扩展]\n\n## 🎯 考公考点\n- 申论角度：[可用主题]\n- 行测考点：[可能出题方向]\n\n## 💡 规范表述\n[2-3句官方标准表述]"
            analysis = call_deepseek(prompt, "你是公考时政分析专家")
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(f"# {title}\n\n")
                f.write(f"**日期**：{today}\n")
                f.write(f"**来源**：{info['source']}\n")
                if info['url']:
                    f.write(f"**原文链接**：{info['url']}\n")
                f.write("\n---\n\n")
                f.write(analysis)
            saved_files.append(filepath)
            print(f"[CRON] 已保存: {filename}")

        # 3. 生成日报
        formatted = "\n".join([f"{i+1}. 【{info['source']}】{title}" for i, (title, info) in enumerate(sorted_news)])
        report_prompt = f"""请分析以下时政新闻，按指定格式输出。

【今日新闻】
{formatted}

请按以下格式输出：

📌 **今日时政核心要点**（3-5条）
- [新闻事件]：一句话概括+考公切入点

📝 **一句话考点总结**
1. 【来源】标题 → 考点方向：[具体考点]

🔑 **高频关键词**
关键词1 · 关键词2 · 关键词3

📖 **申论素材卡**
- 可用主题：[主题]
- 核心案例：[新闻事件]
- 规范表述：「官方表述」

🎯 **行测时政预测**
- 可能出题方向：[方向]

💪 **【今日随堂测验】**

请生成2道与今日时政相关的单选题：

**题目1**
[问题]
A. [选项1]  B. [选项2]  C. [选项3]  D. [选项4]
答案：[_]

**题目2**
[问题]
A. [选项1]  B. [选项2]  C. [选项3]  D. [选项4]
答案：[_]
"""
        report = call_deepseek(report_prompt, "你是一个公考时政分析专家。请严格按照要求的格式输出分析结果。")
        report_path = os.path.join(DAILY_REPORTS_DIR, f"daily_news_{today}.md")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(f"# 📰 {today} 时政日报\n\n")
            f.write(report)
        print(f"[CRON] 日报已生成: {report_path}")

        # 4. 嵌入知识库
        if saved_files:
            try:
                client = get_chroma_client()
                collection = client.get_or_create_collection(
                    name="shizheng_news",
                    metadata={"description": "考公时政新闻知识库"}
                )
                for filepath in saved_files:
                    with open(filepath, "r", encoding="utf-8") as f:
                        content = f.read()
                    chunks = chunk_text(content)
                    if not chunks:
                        continue
                    fname = os.path.basename(filepath).replace(".md", "")
                    date_match = re.match(r'(\d{4}-\d{2}-\d{2})', fname)
                    for j, chunk in enumerate(chunks):
                        try:
                            emb = get_embedding([chunk])
                            if emb:
                                meta = {
                                    "title": fname,
                                    "date": date_match.group(1) if date_match else today,
                                    "source": "新闻归档"
                                }
                                collection.add(
                                    documents=[chunk],
                                    embeddings=emb,
                                    ids=[f"{fname}_chunk_{j}"],
                                    metadatas=[meta]
                                )
                        except Exception as e:
                            print(f"[CRON] 嵌入chunk失败: {e}")
                print(f"[CRON] 知识库已更新，新增{len(saved_files)}篇")
            except Exception as e:
                print(f"[CRON] 知识库嵌入失败: {e}")

        # 5. Bark推送到手机
        bark_key = os.environ.get("BARK_DEVICE_KEY", "CvCwzTHBHUNpm8znAvxsSB")
        if bark_key:
            try:
                bark_url = f"https://api.day.app/{bark_key}"
                bark_body = report[:3500] if len(report) > 3500 else report
                bark_payload = {
                    "title": f"📰 时政日报 {today}",
                    "body": bark_body,
                    "group": "时政日报",
                }
                requests.post(bark_url, json=bark_payload, timeout=10)
                print(f"[CRON] Bark推送成功")
            except Exception as e:
                print(f"[CRON] Bark推送失败: {e}")

        # 6. 自动同步到GitHub（数据持久化）
        try:
            git_sync_to_github(f"auto: 每日爬取 {today}")
        except Exception as e:
            print(f"[CRON] GitHub同步失败: {e}")

        print(f"[CRON] 每日爬取完成")

    # 异步执行，避免超时
    thread = threading.Thread(target=do_crawl)
    thread.start()
    return jsonify({"message": "爬取任务已启动", "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")})


@app.route("/api/rebuild-kb", methods=["POST", "GET"])
def rebuild_knowledge_base():
    """重建知识库：将所有已有新闻归档重新嵌入向量库"""
    secret = request.args.get("secret", "") or request.headers.get("X-Cron-Secret", "")
    if secret != CRON_SECRET:
        return jsonify({"error": "unauthorized"}), 401

    result = _rebuild_kb_from_archives()
    if not result.get("success"):
        status = 404 if result.get("error", "").startswith("无新闻归档") else 500
        return jsonify(result), status
    return jsonify(result)


def _rebuild_kb_from_archives() -> dict:
    """从新闻归档重建向量库，返回结果字典。供API与启动自检共用。"""
    files = glob.glob(os.path.join(NEWS_ARCHIVE_DIR, "*.md"))
    if not files:
        return {"success": False, "error": "无新闻归档文件", "dir": NEWS_ARCHIVE_DIR}

    errors = []
    success_count = 0
    try:
        client = get_chroma_client()
        # 删除旧的再重建
        try:
            client.delete_collection("shizheng_news")
        except Exception:
            pass
        collection = client.get_or_create_collection(
            name="shizheng_news",
            metadata={"description": "考公时政新闻知识库"}
        )

        for filepath in sorted(files):
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()
                if not content.strip():
                    continue
                chunks = chunk_text(content)
                if not chunks:
                    continue
                # 逐条嵌入，避免批量过大
                fname = os.path.basename(filepath).replace(".md", "")
                for j, chunk in enumerate(chunks):
                    try:
                        emb = get_embedding([chunk])
                        if emb:
                            # 提取元数据
                            date_match = re.match(r'(\d{4}-\d{2}-\d{2})', fname)
                            meta = {
                                "title": fname,
                                "date": date_match.group(1) if date_match else "",
                                "source": "新闻归档"
                            }
                            collection.add(
                                documents=[chunk],
                                embeddings=emb,
                                ids=[f"{fname}_chunk_{j}"],
                                metadatas=[meta]
                            )
                    except Exception as e:
                        errors.append(f"{fname}_chunk_{j}: {str(e)[:100]}")
                success_count += 1
            except Exception as e:
                errors.append(f"{os.path.basename(filepath)}: {str(e)[:100]}")

        total = collection.count()
        return {
            "success": True,
            "files_processed": success_count,
            "total_chunks": total,
            "errors": errors[:10]
        }
    except Exception as e:
        return {"success": False, "error": f"知识库重建失败: {str(e)}", "details": errors[:10]}


@app.route("/api/debug/embedding-test")
def debug_embedding():
    """测试嵌入API是否正常"""
    test_text = "习近平总书记关于全面依法治国的重要论述"
    try:
        result = get_embedding([test_text])
        if result:
            return jsonify({"success": True, "dim": len(result[0]), "first_5": result[0][:5]})
        else:
            return jsonify({"success": False, "error": "返回空结果"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/sync-github", methods=["POST", "GET"])
def manual_sync_github():
    """手动触发数据同步到GitHub"""
    secret = request.args.get("secret", "") or request.headers.get("X-Cron-Secret", "")
    if secret != CRON_SECRET:
        return jsonify({"error": "unauthorized"}), 401
    try:
        success = git_sync_to_github("auto: 手动触发同步")
        if success:
            return jsonify({"success": True, "message": "数据已同步到GitHub"})
        else:
            return jsonify({"success": False, "message": "同步失败，请检查GITHUB_TOKEN配置"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ========== 启动 ==========

def _startup_kb_autocheck():
    """启动自检：向量库为空且有归档时，自动后台重建。
    Render等免费平台重新部署会清空文件系统，导致向量库丢失，此处自动恢复。"""
    def _run():
        try:
            client = get_chroma_client()
            try:
                count = client.get_collection("shizheng_news").count()
            except Exception:
                count = 0
            if count > 0:
                print(f"[STARTUP] 知识库已就绪，共 {count} 个文本块，跳过重建")
                return
            files = glob.glob(os.path.join(NEWS_ARCHIVE_DIR, "*.md"))
            if not files:
                print("[STARTUP] 知识库为空且无归档文件，跳过自动重建")
                return
            print(f"[STARTUP] 检测到知识库为空，开始从 {len(files)} 个归档自动重建...")
            result = _rebuild_kb_from_archives()
            if result.get("success"):
                print(f"[STARTUP] 知识库自动重建完成，共 {result.get('total_chunks')} 个文本块")
            else:
                print(f"[STARTUP] 知识库自动重建失败: {result.get('error')}")
        except Exception as e:
            print(f"[STARTUP] 知识库自检异常: {e}")

    import threading
    threading.Thread(target=_run, daemon=True).start()


# 模块加载即触发（兼容 gunicorn，不依赖 __main__）
_startup_kb_autocheck()


if __name__ == "__main__":
    print("=" * 50)
    print("  考公时政知识库系统")
    port = int(os.environ.get("PORT", 5678))
    print(f"  访问地址: http://localhost:{port}")
    print("=" * 50)
    app.run(host="0.0.0.0", port=port, debug=os.environ.get("FLASK_DEBUG", "1") == "1")

